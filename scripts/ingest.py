#!/usr/bin/env python3
"""
HCT Library AI Agent — Document Ingestion Pipeline
====================================================
Extracts and chunks text from PDF, DOCX, and HTML files.
Outputs structured JSON chunks ready for embedding.

Usage:
    python scripts/ingest.py --source documents/ --output chunks/
    python scripts/ingest.py --file documents/library_guide.pdf
    python scripts/ingest.py --source documents/ --output chunks/ --verbose

Requirements:
    pip install pymupdf python-docx beautifulsoup4 trafilatura tiktoken

Output format (chunks/YYYYMMDD_HHMMSS.jsonl):
    {"id": "...", "text": "...", "source": "filename.pdf", "category": "reference",
     "chunk_index": 0, "token_count": 387, "metadata": {...}}
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path

# ── CONFIGURATION ──────────────────────────────────────────────────────────────
CHUNK_SIZE_TOKENS = 400   # Target tokens per chunk
CHUNK_OVERLAP_TOKENS = 50  # Overlap between consecutive chunks
MIN_CHUNK_TOKENS = 50      # Discard chunks shorter than this
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.html', '.htm', '.txt', '.json'}

# ── HELPERS ───────────────────────────────────────────────────────────────────
def count_tokens(text: str) -> int:
    """Approximate token count (4 chars ≈ 1 token for English)."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except ImportError:
        return len(text) // 4


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE_TOKENS,
               overlap: int = CHUNK_OVERLAP_TOKENS) -> list[str]:
    """Split text into overlapping chunks by approximate token count."""
    words = text.split()
    chars_per_token = 4
    chunk_words = chunk_size * chars_per_token // 5  # approximate words per chunk
    overlap_words = overlap * chars_per_token // 5

    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_words, len(words))
        chunk = ' '.join(words[start:end]).strip()
        if count_tokens(chunk) >= MIN_CHUNK_TOKENS:
            chunks.append(chunk)
        if end >= len(words):
            break
        start = end - overlap_words
    return chunks


def make_chunk_id(source: str, index: int) -> str:
    """Generate a deterministic chunk ID."""
    raw = f"{source}::{index}"
    return hashlib.md5(raw.encode()).hexdigest()[:16]


def clean_text(text: str) -> str:
    """Normalise whitespace and remove junk characters."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E -￿]', '', text)
    return text.strip()


# ── EXTRACTORS ────────────────────────────────────────────────────────────────
def extract_pdf(filepath: Path) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("  [!] PyMuPDF not installed. Run: pip install pymupdf", file=sys.stderr)
        return ""
    doc = fitz.open(str(filepath))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return clean_text('\n'.join(pages))


def extract_docx(filepath: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        print("  [!] python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return ""
    doc = Document(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return clean_text('\n'.join(paragraphs))


def extract_html(filepath: Path) -> str:
    """Extract main content from HTML using trafilatura with BeautifulSoup fallback."""
    html_content = filepath.read_text(encoding='utf-8', errors='ignore')
    try:
        import trafilatura
        text = trafilatura.extract(html_content)
        if text:
            return clean_text(text)
    except ImportError:
        pass
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        return clean_text(soup.get_text(separator='\n'))
    except ImportError:
        print("  [!] BeautifulSoup4 not installed. Run: pip install beautifulsoup4", file=sys.stderr)
        return ""


def extract_txt(filepath: Path) -> str:
    return clean_text(filepath.read_text(encoding='utf-8', errors='ignore'))


def extract_faq_json(filepath: Path) -> list[dict]:
    """Special handler for the FAQ JSON format (returns pre-formed chunks)."""
    data = json.loads(filepath.read_text(encoding='utf-8'))
    faqs = data if isinstance(data, list) else data.get('faqs', [])
    chunks = []
    for i, item in enumerate(faqs):
        question = item.get('question', '')
        answer = item.get('answer', '')
        tags = ', '.join(item.get('tags', []))
        text = f"Q: {question}\nA: {answer}"
        if tags:
            text += f"\nTopics: {tags}"
        chunks.append({
            'text': clean_text(text),
            'metadata': {'category': item.get('category', 'general'), 'tags': item.get('tags', [])}
        })
    return chunks


# ── MAIN INGESTION ────────────────────────────────────────────────────────────
def ingest_file(filepath: Path, category: str = 'general', verbose: bool = False) -> list[dict]:
    """Ingest a single file and return a list of chunk dicts."""
    ext = filepath.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        if verbose:
            print(f"  [skip] Unsupported extension: {filepath.name}")
        return []

    if verbose:
        print(f"  Processing: {filepath.name}")

    # Special case: FAQ JSON
    if ext == '.json':
        pre_chunks = extract_faq_json(filepath)
        result = []
        for i, pc in enumerate(pre_chunks):
            result.append({
                'id': make_chunk_id(filepath.name, i),
                'text': pc['text'],
                'source': filepath.name,
                'category': pc.get('metadata', {}).get('category', category),
                'chunk_index': i,
                'token_count': count_tokens(pc['text']),
                'metadata': pc.get('metadata', {}),
            })
        return result

    # Text-based files
    if ext == '.pdf':
        text = extract_pdf(filepath)
    elif ext == '.docx':
        text = extract_docx(filepath)
    elif ext in ('.html', '.htm'):
        text = extract_html(filepath)
    else:
        text = extract_txt(filepath)

    if not text:
        if verbose:
            print(f"  [warn] No text extracted from {filepath.name}")
        return []

    raw_chunks = chunk_text(text)
    result = []
    for i, chunk in enumerate(raw_chunks):
        result.append({
            'id': make_chunk_id(filepath.name, i),
            'text': chunk,
            'source': filepath.name,
            'category': category,
            'chunk_index': i,
            'token_count': count_tokens(chunk),
            'metadata': {
                'filepath': str(filepath),
                'total_chunks': len(raw_chunks),
            },
        })
    return result


def ingest_directory(source_dir: Path, verbose: bool = False) -> list[dict]:
    """Recursively ingest all supported files in a directory."""
    all_chunks = []
    for filepath in sorted(source_dir.rglob('*')):
        if filepath.is_file() and filepath.suffix.lower() in SUPPORTED_EXTENSIONS:
            # Infer category from parent folder name
            category = filepath.parent.name if filepath.parent != source_dir else 'general'
            chunks = ingest_file(filepath, category=category, verbose=verbose)
            all_chunks.extend(chunks)
            if verbose and chunks:
                print(f"    → {len(chunks)} chunks")
    return all_chunks


# ── CLI ───────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='HCT Library Document Ingestion Pipeline')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--source', type=Path, help='Directory of documents to ingest')
    group.add_argument('--file', type=Path, help='Single file to ingest')
    parser.add_argument('--output', type=Path, default=Path('chunks'),
                        help='Output directory for JSONL chunks (default: chunks/)')
    parser.add_argument('--category', default='general', help='Category label for single-file mode')
    parser.add_argument('--verbose', action='store_true', help='Verbose output')
    args = parser.parse_args()

    print("HCT Library Ingestion Pipeline")
    print("=" * 40)
    start = time.time()

    # Collect chunks
    if args.file:
        if not args.file.exists():
            print(f"Error: File not found: {args.file}", file=sys.stderr)
            sys.exit(1)
        chunks = ingest_file(args.file, category=args.category, verbose=True)
    else:
        if not args.source.exists():
            print(f"Error: Directory not found: {args.source}", file=sys.stderr)
            sys.exit(1)
        print(f"Source: {args.source}")
        chunks = ingest_directory(args.source, verbose=args.verbose)

    if not chunks:
        print("No chunks produced. Check input files and dependencies.")
        sys.exit(0)

    # Write output
    args.output.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    out_path = args.output / f'chunks_{timestamp}.jsonl'
    with open(out_path, 'w', encoding='utf-8') as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + '\n')

    elapsed = time.time() - start
    print(f"\nDone in {elapsed:.1f}s")
    print(f"Chunks produced : {len(chunks)}")
    print(f"Output file     : {out_path}")
    print(f"\nNext step: python scripts/embed.py --input {out_path}")


if __name__ == '__main__':
    main()
